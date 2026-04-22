#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成API文档的Word文档
"""
import os
import sys
from docx import Document
from docx.shared import Inches

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import create_app
from app.routes import *

app = create_app('default')

# API路由映射
API_ROUTES = {
    '用户模块': {
        'routes': [
            ('/users/send-email-code', 'POST', '发送邮箱验证码', '否'),
            ('/users/send-sms', 'POST', '发送短信验证码', '否'),
            ('/users/register', 'POST', '用户注册（老人或家属）', '是'),
            ('/users/login', 'POST', '用户登录', '是'),
            ('/users/profile', 'GET', '获取用户资料', '是'),
            ('/users/profile', 'PUT', '更新用户资料', '是'),
            ('/users/change-password', 'POST', '修改密码', '否'),
            ('/users/elder/list', 'GET', '获取老人列表', '是'),
            ('/users/binding-elder', 'GET', '获取家属绑定的老人信息', '是'),
            ('/users/binding-elder', 'POST', '家属绑定老人', '是'),
            ('/users/binding-elder', 'DELETE', '家属解绑老人', '是'),
            ('/users/workers', 'GET', '获取护工列表', '是'),
            ('/users/elder', 'POST', '管理员创建老人账号', '是'),
            ('/users/<int:user_id>', 'PUT', '管理员更新用户信息', '是'),
            ('/users/<int:user_id>', 'DELETE', '管理员删除用户（仅限老人）', '是'),
        ]
    },
    '订单模块': {
        'routes': [
            ('/orders/summary', 'GET', '获取订单统计信息', '是'),
            ('/orders/', 'GET', '获取订单列表', '是'),
            ('/orders/<int:order_id>', 'GET', '获取订单详情', '是'),
            ('/orders/', 'POST', '创建订单', '是'),
            ('/orders/<int:order_id>', 'PUT', '更新订单信息', '是'),
            ('/orders/<int:order_id>', 'DELETE', '删除订单', '是'),
        ]
    },
    '护理记录模块': {
        'routes': [
            ('/nursing/records', 'POST', '创建护理记录', '是'),
            ('/nursing/records', 'GET', '获取护理记录列表', '是'),
            ('/nursing/records/<int:record_id>', 'GET', '获取护理记录详情', '否'),
            ('/nursing/records/<int:record_id>', 'PUT', '更新护理记录', '否'),
            ('/nursing/records/<int:record_id>/complete', 'POST', '标记护理记录为已完成', '否'),
            ('/nursing/types', 'GET', '获取护理类型列表', '是'),
        ]
    },
    '用药提醒模块': {
        'routes': [
            ('/medication/reminders', 'GET', '获取用药提醒列表', '是'),
            ('/medication/reminders', 'POST', '创建用药提醒', '是'),
            ('/medication/reminders/<int:reminder_id>/complete', 'PUT', '标记用药提醒为已完成', '是'),
            ('/medication/reminders/<int:reminder_id>/uncomplete', 'PUT', '标记用药提醒为未完成', '是'),
            ('/medication/reminders/<int:reminder_id>', 'DELETE', '删除用药提醒', '是'),
        ]
    },
    '位置模块': {
        'routes': [
            ('/location/update', 'POST', '更新位置信息', '否'),
            ('/location/current', 'GET', '获取当前位置', '否'),
            ('/location/history', 'GET', '获取位置历史记录', '否'),
            ('/location/elder/<int:elder_id>/current', 'GET', '获取老人当前位置', '否'),
            ('/location/elder/<int:elder_id>/history', 'GET', '获取老人位置历史记录', '否'),
        ]
    },
    '紧急呼叫模块': {
        'routes': [
            ('/emergency/call', 'POST', '发起紧急呼叫', '是'),
            ('/emergency/calls', 'GET', '获取紧急呼叫列表', '是'),
            ('/emergency/calls/<int:call_id>/respond', 'POST', '响应紧急呼叫', '是'),
            ('/emergency/calls/<int:call_id>/complete', 'POST', '完成紧急呼叫', '是'),
        ]
    },
    '打卡模块': {
        'routes': [
            ('/checkin/', 'POST', '执行打卡', '是'),
            ('/checkin/history', 'GET', '获取打卡历史', '是'),
            ('/checkin/stats', 'GET', '获取打卡统计', '否'),
            ('/checkin/admin/history', 'GET', '获取所有用户的打卡历史', '是'),
        ]
    },
    '护理计划模块': {
        'routes': [
            ('/care/plans', 'POST', '创建护理计划', '是'),
            ('/care/plans', 'GET', '获取护理计划列表', '是'),
            ('/care/plans/<int:plan_id>', 'GET', '获取护理计划详情', '否'),
            ('/care/plans/<int:plan_id>', 'PUT', '更新护理计划', '否'),
            ('/care/tasks', 'POST', '创建护理任务', '否'),
            ('/care/tasks/<int:task_id>/complete', 'POST', '标记护理任务为已完成', '否'),
            ('/care/tasks/today', 'GET', '获取今日护理任务', '否'),
            ('/care/plans/<int:plan_id>/tasks', 'POST', '为护理计划添加任务', '否'),
            ('/care/plans/today', 'GET', '获取今日护理计划', '是'),
        ]
    },
    '通知模块': {
        'routes': [
            ('/notifications/', 'POST', '创建通知', '否'),
            ('/notifications/', 'GET', '获取通知列表', '是'),
            ('/notifications/<int:notification_id>/read', 'POST', '标记通知为已读', '否'),
            ('/notifications/read-all', 'POST', '标记所有通知为已读', '否'),
            ('/notifications/unread-count', 'GET', '获取未读通知数量', '是'),
            ('/notifications/broadcast', 'POST', '广播通知', '否'),
            ('/notifications/types', 'GET', '获取通知类型', '否'),
            ('/notifications/users', 'GET', '获取用户列表（用于发送通知）', '是'),
            ('/notifications/create', 'POST', '创建通知', '是'),
        ]
    },
    '大数据模块': {
        'routes': [
            ('/bigdata/health-analysis', 'GET', '健康分析', '否'),
            ('/bigdata/elder-health-report', 'GET', '老人健康报告', '否'),
            ('/bigdata/health-prediction', 'GET', '健康预测', '否'),
            ('/bigdata/nursing-efficiency', 'GET', '护理效率分析', '否'),
            ('/bigdata/care-plan-analysis', 'GET', '护理计划分析', '否'),
            ('/bigdata/risk-assessment', 'GET', '风险评估', '否'),
            ('/bigdata/activity-patterns', 'GET', '活动模式分析', '否'),
            ('/bigdata/comparison', 'GET', '健康数据对比', '否'),
            ('/bigdata/anomaly-alerts', 'GET', '异常提醒', '否'),
            ('/bigdata/health-insights', 'GET', '健康洞察', '否'),
        ]
    },
    '支付模块': {
        'routes': [
            ('/payment/create', 'POST', '创建支付', '否'),
            ('/payment/query/<order_no>', 'GET', '查询支付状态', '否'),
            ('/payment/notify', 'POST/GET', '支付通知', '否'),
            ('/payment/refund', 'POST', '退款', '否'),
            ('/payment/cancel/<int:order_id>', 'POST', '取消支付', '否'),
            ('/payment/wechat/auth', 'GET', '微信授权', '否'),
            ('/payment/config', 'GET', '获取支付配置', '否'),
            ('/payment/mock', 'GET/POST', '模拟支付', '否'),
        ]
    },
    'AI模块': {
        'routes': [
            ('/ai/chat', 'POST', 'AI聊天', '是'),
            ('/ai/health-check', 'GET', '健康检查', '否'),
            ('/ai/models', 'GET', '获取AI模型列表', '否'),
            ('/ai/config', 'GET', '获取AI配置', '否'),
            ('/ai/config', 'PUT', '更新AI配置', '否'),
        ]
    },
    '服务模块': {
        'routes': [
            ('/service/', 'GET', '获取服务列表', '是'),
            ('/service/<int:service_id>', 'GET', '获取服务详情', '否'),
            ('/service/categories', 'GET', '获取服务分类', '否'),
            ('/service/', 'POST', '创建服务', '否'),
            ('/service/<int:service_id>', 'PUT', '更新服务', '否'),
            ('/service/<int:service_id>', 'DELETE', '删除服务', '否'),
        ]
    },
    '统计模块': {
        'routes': [
            ('/statistics/dashboard', 'GET', '获取仪表盘统计数据', '是'),
            ('/statistics/nursing-summary', 'GET', '获取护理统计摘要', '否'),
            ('/statistics/health-trend', 'GET', '获取健康趋势', '否'),
        ]
    },
    '评价模块': {
        'routes': [
            ('/evaluation/worker', 'POST', '创建护工评价', '是'),
            ('/evaluation/worker', 'GET', '获取护工评价列表', '是'),
            ('/evaluation/worker/<int:evaluation_id>', 'GET', '获取评价详情', '否'),
            ('/evaluation/worker/<int:evaluation_id>/reply', 'POST', '护工回复评价', '否'),
            ('/evaluation/worker/stats/<int:worker_id>', 'GET', '获取护工评分统计', '是'),
        ]
    },
    '健康模块': {
        'routes': [
            ('/health/metrics/latest/<int:elder_id>', 'GET', '获取老人最新健康数据', '是'),
        ]
    },
    '电子围栏模块': {
        'routes': [
            ('/geofence/', 'GET', '获取电子围栏列表', '否'),
            ('/geofence/', 'POST', '创建电子围栏', '否'),
            ('/geofence/<int:fence_id>', 'PUT', '更新电子围栏', '否'),
            ('/geofence/<int:fence_id>', 'DELETE', '删除电子围栏', '否'),
            ('/geofence/alerts', 'GET', '获取电子围栏告警', '否'),
        ]
    },
    '系统模块': {
        'routes': [
            ('/system/info', 'GET', '获取系统信息', '否'),
            ('/system/status', 'GET', '获取系统状态', '否'),
        ]
    },
}

def generate_api_doc():
    """生成API文档"""
    # 创建文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('智能护理系统 API 文档', 0)
    
    # 添加版本信息
    doc.add_paragraph('版本: 1.0.0')
    doc.add_paragraph('生成日期: 2026-04-22')
    doc.add_paragraph('')
    
    # 添加目录
    doc.add_heading('目录', level=1)
    for module_name in API_ROUTES:
        doc.add_paragraph(f'1. {module_name}', style='List Number')
    doc.add_paragraph('')
    
    # 遍历每个模块
    for module_index, module_name in enumerate(API_ROUTES, start=1):
        # 添加模块标题
        doc.add_heading(f'{module_index}. {module_name}', level=1)
        
        # 获取模块的路由
        routes = API_ROUTES[module_name]['routes']
        
        # 创建表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'API路径'
        hdr_cells[1].text = '方法'
        hdr_cells[2].text = '功能描述'
        hdr_cells[3].text = '是否被使用'
        
        # 填充表格数据
        for route in routes:
            row_cells = table.add_row().cells
            row_cells[0].text = route[0]
            row_cells[1].text = route[1]
            row_cells[2].text = route[2]
            row_cells[3].text = route[3]
        
        doc.add_paragraph('')
    
    # 添加总结
    doc.add_heading('总结', level=1)
    doc.add_paragraph('已使用的API：大部分核心功能的API都已被使用，包括用户登录注册、订单管理、护理记录、用药提醒、紧急呼叫、打卡、通知等。')
    doc.add_paragraph('未使用的API：主要是一些高级功能或管理功能的API，如大数据分析、支付、AI健康检查、电子围栏等。')
    doc.add_paragraph('需要注意的API：部分API虽然存在，但前端代码中没有直接调用，可能需要根据实际需求进行集成。')
    
    # 保存文档
    doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'api_document.docx')
    doc.save(doc_path)
    print(f'API文档已生成：{doc_path}')

if __name__ == '__main__':
    generate_api_doc()
