# -*- coding: utf-8 -*-
"""
生成智慧护理平台架构设计与详细设计说明书 Word文档
严格按照用户提供的目录结构
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), edge_data.get('sz', '4'))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading_centered(doc, text, level):
    """添加居中标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def add_table_with_header(doc, headers, rows, header_color='4472C4'):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    return table


def create_document():
    """创建完整的架构设计文档"""
    doc = Document()
    
    # ========== 封面页 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run('互联网+智慧护理移动护理平台')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('架构设计与详细设计说明书')
    subtitle_run.font.size = Pt(22)
    subtitle_run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 项目信息表
    info_table = doc.add_table(rows=9, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ['卷  号', '', '卷内编号', ''],
        ['密  级', '', '', ''],
        ['项目编号', 'HD20221101SR005', '文档编号', 'HD20221101SR005'],
        ['', '', '', ''],
        ['项目名称', '互联网+智慧护理移动护理平台APP', '', ''],
        ['项目承担部门', '坤坤快跑技术有限公司', '', ''],
        ['撰  写  人', '李鹏辉', '完成日期', '2026-4-15'],
        ['评审负责人', '江雨泽', '评审日期', '2026-4-15'],
        ['承建单位', '四川华迪信息技术有限公司', '', '']
    ]
    
    for i, row_data in enumerate(info_data):
        for j, cell_text in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = cell_text
            if j % 2 == 0:
                set_cell_shading(cell, 'E8E8E8')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_heading_centered(doc, '目  录', 1)
    doc.add_paragraph()
    
    toc_items = [
        ('1. 简介', '1'),
        ('1.1 目的', '1'),
        ('1.2 范围', '1'),
        ('1.3 定义、首字母缩写词和缩略语', '1'),
        ('1.4 参考资料', '1'),
        ('1.5 概述', '1'),
        ('2. 架构设计', '1'),
        ('2.1 数据处理架构设计', '1'),
        ('2.2 数据分析和挖掘架构设计', '1'),
        ('2.3 数据安全架构设计', '1'),
        ('3. 构架表示方式', '1'),
        ('4. 构架目标和约束', '1'),
        ('5. 关键用例视图', '2'),
        ('5.1 老人管理', '2'),
        ('5.2 护理人员管理', '2'),
        ('5.3 护理记录管理', '2'),
        ('5.4 订单管理', '2'),
        ('5.5 健康数据管理', '2'),
        ('5.6 护理计划管理', '3'),
        ('6. 层次结构', '3'),
        ('7. 逻辑视图', '3'),
        ('7.1 概述', '3'),
        ('7.2 用户服务层', '4'),
        ('7.3 业务逻辑层', '5'),
        ('7.4 数据服务层', '6'),
        ('8. 部署视图', '6'),
        ('9. 接口设计', '7'),
        ('10. 系统的组织结构', '7'),
        ('11. 界面设计要求', '7'),
        ('12. 详细设计', '8'),
        ('12.1 系统模块划分', '8'),
        ('12.2 模块详细设计', '8'),
        ('12.3 接口设计', '8'),
        ('12.4 算法与模型设计', '8'),
        ('12.5 系统部署与运维设计', '8'),
        ('12.6 登入注册界面的设计说明', '8'),
        ('12.7 护理计划的设计说明', '9'),
        ('12.8 健康数据的设计说明', '10'),
        ('12.9 大数据分析的设计说明', '11'),
        ('13. 模块相互关系表', '12'),
    ]
    
    for title_text, page_num in toc_items:
        p = doc.add_paragraph()
        p.add_run(title_text)
        tab_count = (60 - len(title_text)) // 2
        p.add_run('\t' * tab_count + page_num)
    
    doc.add_page_break()
    
    # ========== 1. 简介 ==========
    doc.add_heading('1. 简介', level=1)
    
    doc.add_heading('1.1 目的', level=2)
    doc.add_paragraph(
        '本文档旨在详细描述互联网+智慧护理移动护理平台的技术架构和设计细节，'
        '为开发团队提供完整的技术参考，确保系统开发的一致性和可维护性。'
        '文档涵盖系统架构、后端设计、前端设计、数据库设计、API接口设计等各个方面，'
        '为系统的实现、测试、部署和维护提供技术依据。'
    )
    
    doc.add_heading('1.2 范围', level=2)
    doc.add_paragraph(
        '本设计文档适用于智慧护理平台的完整技术实现，包括以下内容：\n\n'
        '• 前端应用：Vue 3 + TypeScript + Element Plus\n'
        '• 后端服务：Flask + Python\n'
        '• 数据库：SQLite\n'
        '• AI服务：科大讯飞语音识别\n'
        '• 实时通信：WebSocket\n'
        '• 数据分析：ECharts可视化'
    )
    
    doc.add_heading('1.3 定义、首字母缩写词和缩略语', level=2)
    
    glossary_table = doc.add_table(rows=12, cols=2)
    glossary_table.style = 'Table Grid'
    glossary_data = [
        ['术语/缩写', '定义'],
        ['API', '应用程序编程接口（Application Programming Interface）'],
        ['JWT', 'JSON Web Token，用于用户身份认证'],
        ['CRUD', '创建、读取、更新、删除（Create, Read, Update, Delete）'],
        ['RESTful', '一种API设计风格，遵循HTTP协议语义'],
        ['ORM', '对象关系映射（Object-Relational Mapping）'],
        ['Vue', '渐进式JavaScript框架'],
        ['Flask', '轻量级Python Web框架'],
        ['SQLite', '轻量级关系型数据库'],
        ['ECharts', '百度开源的数据可视化图表库'],
        ['user_type', '用户类型：1-老人，2-护理人员，3-管理员，4-家属'],
        ['护理类型', '1-日常照护，2-医疗护理，3-康复训练，4-心理疏导，5-饮食护理，6-清洁护理，7-安全护理']
    ]
    for i, row_data in enumerate(glossary_data):
        for j, cell_text in enumerate(row_data):
            cell = glossary_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    doc.add_heading('1.4 参考资料', level=2)
    doc.add_paragraph(
        '• 《软件工程文档规范》（GB/T 8567-2006）\n'
        '• 《智慧护理系统需求规格说明书》\n'
        '• Flask官方文档：https://flask.palletsprojects.com/\n'
        '• Vue 3官方文档：https://vuejs.org/\n'
        '• Element Plus组件库文档：https://element-plus.org/'
    )
    
    doc.add_heading('1.5 概述', level=2)
    doc.add_paragraph(
        '互联网+智慧护理移动护理平台是一个面向养老护理行业的综合性管理平台。'
        '系统采用前后端分离架构，前端使用Vue 3构建单页面应用，后端使用Flask提供RESTful API服务。'
        '平台支持多种用户角色（老人、护理人员、管理员、家属），提供护理服务预约、'
        '健康数据管理、护理计划制定、大数据分析等功能。'
    )
    
    doc.add_page_break()
    
    # ========== 2. 架构设计 ==========
    doc.add_heading('2. 架构设计', level=1)
    
    doc.add_heading('2.1 数据处理架构设计', level=2)
    doc.add_paragraph(
        '系统数据处理采用分层架构模式，主要包括以下几个层次：\n\n'
        '（1）数据采集层\n'
        '负责从各种来源收集数据，包括用户输入、健康设备采集、第三方系统对接等。'
        '支持多种数据格式（JSON、XML、CSV等），并对数据进行初步校验和格式化。\n\n'
        '（2）数据存储层\n'
        '使用SQLite数据库进行持久化存储，采用ORM技术（SQLAlchemy）进行数据库操作。'
        '数据库包含用户表、健康指标表、护理记录表、订单表、服务表等核心数据表。\n\n'
        '（3）数据处理层\n'
        '后端业务逻辑层负责处理业务规则、数据验证、计算统计等。'
        '使用Python的statistics库进行数据统计分析，支持均值、中位数、标准差等计算。\n\n'
        '（4）数据展示层\n'
        '前端使用ECharts进行数据可视化，支持折线图、柱状图、饼图等多种图表类型。'
        '数据以JSON格式通过RESTful API进行传输。'
    )
    
    doc.add_heading('2.2 数据分析和挖掘架构设计', level=2)
    doc.add_paragraph(
        '系统数据分析模块采用以下架构设计：\n\n'
        '（1）健康指标分析\n'
        '• 统计指标计算：最小值、最大值、平均值、中位数、标准差、方差\n'
        '• 异常检测：基于健康标准范围判断指标是否异常\n'
        '• 趋势分析：对比近期数据与历史数据，判断指标变化趋势\n\n'
        '（2）风险评估\n'
        '• 异常率计算：统计各项指标的异常比例\n'
        '• 风险分级：高风险（异常率>50%）、中风险（异常率>20%）、低风险\n'
        '• 风险因素识别：分析可能导致健康风险的因素\n\n'
        '（3）健康预测\n'
        '• 线性回归模型：基于历史数据预测未来趋势\n'
        '• 预测周期：支持自定义预测天数\n'
        '• 置信度评估：提供预测结果的置信区间\n\n'
        '（4）智能洞察生成\n'
        '• 数据完整度分析：评估数据采集的完整性\n'
        '• 测量频率分析：统计各指标的测量频次\n'
        '• AI建议生成：基于分析结果提供健康建议'
    )
    
    doc.add_heading('2.3 数据安全架构设计', level=2)
    doc.add_paragraph(
        '（1）用户认证与授权\n'
        '• JWT Token认证：用户登录后获取Token，后续请求携带Token进行身份验证\n'
        '• Token有效期：支持配置Token过期时间\n'
        '• 角色权限控制：根据用户类型（老人、护理人员、管理员、家属）控制访问权限\n\n'
        '（2）数据加密\n'
        '• 密码加密：使用werkzeug的generate_password_hash进行SHA-256加密\n'
        '• 敏感数据传输：使用HTTPS协议加密传输\n'
        '• 敏感数据存储：对身份证号、联系电话等进行脱敏处理\n\n'
        '（3）输入验证\n'
        '• 后端参数校验：使用Flask-WTF进行表单验证\n'
        '• SQL注入防护：使用ORM避免直接SQL拼接\n'
        '• XSS防护：前端对用户输入进行HTML转义\n\n'
        '（4）接口保护\n'
        '• @require_token装饰器：保护所有需要认证的API接口\n'
        '• CORS配置：控制跨域访问\n'
        '• 请求限流：防止恶意请求'
    )
    
    doc.add_page_break()
    
    # ========== 3. 构架表示方式 ==========
    doc.add_heading('3. 构架表示方式', level=1)
    doc.add_paragraph(
        '系统架构采用以下表示方式：\n\n'
        '（1）分层架构图\n'
        '使用分层架构展示系统的纵向组织结构，包括表现层、业务逻辑层、数据访问层。\n\n'
        '（2）模块结构图\n'
        '使用模块图展示系统的横向功能划分，包括用户管理模块、护理管理模块、健康管理模块、订单管理模块等。\n\n'
        '（3）时序图\n'
        '使用时序图展示关键业务流程的时序交互，如用户登录流程、订单创建流程等。\n\n'
        '（4）ER图\n'
        '使用实体关系图展示数据库表之间的关联关系。\n\n'
        '（5）部署图\n'
        '使用部署图展示系统的物理部署架构，包括前端、后端、数据库的部署方式。'
    )
    
    # ========== 4. 构架目标和约束 ==========
    doc.add_heading('4. 构架目标和约束', level=1)
    doc.add_paragraph(
        '（1）架构目标\n'
        '• 可扩展性：支持功能模块的灵活扩展和裁剪\n'
        '• 可维护性：代码结构清晰，便于维护和升级\n'
        '• 高可用性：支持7x24小时不间断服务\n'
        '• 安全性：确保用户数据安全和隐私保护\n'
        '• 性能优化：支持高并发访问，响应时间<3秒\n\n'
        '（2）技术约束\n'
        '• 前端：必须使用Vue 3 + TypeScript\n'
        '• 后端：必须使用Flask + Python\n'
        '• 数据库：SQLite（开发环境）/ PostgreSQL（生产环境）\n'
        '• 必须支持移动端响应式布局\n\n'
        '（3）业务约束\n'
        '• 必须支持多种用户角色\n'
        '• 必须符合养老护理行业规范\n'
        '• 必须支持数据导出功能'
    )
    
    doc.add_page_break()
    
    # ========== 5. 关键用例视图 ==========
    doc.add_heading('5. 关键用例视图', level=1)
    
    doc.add_heading('5.1 老人管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '管理平台内的老人用户信息，包括基本信息、健康档案、护理需求等。\n\n'
        '（2）主要用例\n'
        '• 查看老人列表\n'
        '• 添加老人信息\n'
        '• 编辑老人资料\n'
        '• 查看老人健康档案\n'
        '• 关联家属信息\n\n'
        '（3）代码实现\n'
        'API路径：/api/users/elder/list\n'
        '模型：User (user_type=1)\n'
        '关键字段：id, username, phone, real_name, gender, age, address, avatar'
    )
    
    doc.add_heading('5.2 护理人员管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '管理平台内的护理人员信息，包括执业信息、服务评价、工作统计等。\n\n'
        '（2）主要用例\n'
        '• 查看护理人员列表\n'
        '• 查看护理人员详情\n'
        '• 查看护理人员评价\n'
        '• 查看护理人员工作统计\n\n'
        '（3）代码实现\n'
        '模型：User (user_type=2)\n'
        '关联表：NursingRecord, Order'
    )
    
    doc.add_heading('5.3 护理记录管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '记录和管理老人的护理服务过程，包括护理类型、开始/结束时间、执行人员等。\n\n'
        '（2）护理类型\n'
        '• 1-日常照护\n'
        '• 2-医疗护理\n'
        '• 3-康复训练\n'
        '• 4-心理疏导\n'
        '• 5-饮食护理\n'
        '• 6-清洁护理\n'
        '• 7-安全护理\n\n'
        '（3）代码实现\n'
        'API路径：/api/nursing/records\n'
        '模型：NursingRecord\n'
        '主要字段：elder_id, nursing_type, description, start_time, end_time, staff_id, status'
    )
    
    doc.add_heading('5.4 订单管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '管理护理服务订单，包括下单、支付、完成、评价等全流程。\n\n'
        '（2）订单状态\n'
        '• 0-已取消\n'
        '• 1-待支付\n'
        '• 2-待服务\n'
        '• 3-服务中\n'
        '• 4-已完成\n'
        '• 5-已退款\n\n'
        '（3）代码实现\n'
        'API路径：/api/orders/\n'
        '模型：Order\n'
        '主要字段：order_no, user_id, service_id, elder_id, nurse_id, total_amount, status'
    )
    
    doc.add_heading('5.5 健康数据管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '记录和管理老人的健康指标数据，支持多种类型的健康监测。\n\n'
        '（2）指标类型\n'
        '• 1-体温（正常范围：36.0-37.3°C）\n'
        '• 2-血压-收缩压\n'
        '• 3-血压-舒张压\n'
        '• 4-心率（正常范围：60-100 BPM）\n'
        '• 5-血氧（正常范围：95-100%）\n'
        '• 6-血糖（正常范围：3.9-6.1 mmol/L）\n'
        '• 7-体重\n'
        '• 8-身高\n'
        '• 9-睡眠时长\n'
        '• 10-今日步数\n\n'
        '（3）代码实现\n'
        'API路径：/api/health/metrics\n'
        '模型：HealthMetric\n'
        '主要字段：elder_id, metric_type, metric_value, unit, recorded_at, notes'
    )
    
    doc.add_heading('5.6 护理计划管理', level=2)
    doc.add_paragraph(
        '（1）功能描述\n'
        '制定和管理老人的个性化护理计划，支持任务分解、进度跟踪、效果评估。\n\n'
        '（2）计划状态\n'
        '• 1-执行中\n'
        '• 2-已完成\n'
        '• 3-已暂停\n\n'
        '（3）任务状态\n'
        '• 1-待执行\n'
        '• 2-已完成\n'
        '• 3-已取消\n\n'
        '（4）代码实现\n'
        'API路径：/api/care/plans, /api/care/tasks\n'
        '模型：CarePlan, CareTask\n'
        '关系：一个护理计划包含多个护理任务'
    )
    
    doc.add_page_break()
    
    # ========== 6. 层次结构 ==========
    doc.add_heading('6. 层次结构', level=1)
    doc.add_paragraph('系统采用四层架构设计：')
    
    layer_table = doc.add_table(rows=5, cols=3)
    layer_table.style = 'Table Grid'
    layer_data = [
        ['层次', '说明', '技术实现'],
        ['表现层', '用户界面展示和交互', 'Vue 3 + Element Plus + ECharts'],
        ['业务逻辑层', '业务规则处理和流程控制', 'Flask + Python'],
        ['数据访问层', '数据库操作和数据存储', 'SQLAlchemy ORM + SQLite'],
        ['基础设施层', '服务器、网络、缓存等基础服务', 'Nginx + Redis + HTTPS']
    ]
    for i, row_data in enumerate(layer_data):
        for j, cell_text in enumerate(row_data):
            cell = layer_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_page_break()
    
    # ========== 7. 逻辑视图 ==========
    doc.add_heading('7. 逻辑视图', level=1)
    
    doc.add_heading('7.1 概述', level=2)
    doc.add_paragraph(
        '系统的逻辑视图展示了系统的功能组织结构，按照服务类型进行模块划分。'
        '主要包含以下服务模块：\n\n'
        '（1）用户服务模块（user_bp）\n'
        '负责用户注册、登录、个人信息管理、老人列表查询等功能。\n\n'
        '（2）护理服务模块（nursing_bp）\n'
        '负责护理记录的创建、查询、更新、删除等CRUD操作。\n\n'
        '（3）健康管理模块（health_bp）\n'
        '负责健康指标的记录、查询、统计等健康管理功能。\n\n'
        '（4）护理计划模块（care_bp）\n'
        '负责护理计划的制定、任务分配、进度跟踪等功能。\n\n'
        '（5）订单服务模块（order_bp）\n'
        '负责服务订单的创建、支付、完成等订单管理功能。\n\n'
        '（6）服务管理模块（service_bp）\n'
        '负责服务项目的配置、分类、推荐等功能。\n\n'
        '（7）数据分析模块（bigdata_bp）\n'
        '负责健康数据的统计、分析、预测、可视化等功能。\n\n'
        '（8）AI服务模块（ai_bp）\n'
        '负责AI健康助手的对话、语音识别、健康建议等功能。'
    )
    
    doc.add_heading('7.2 用户服务层', level=2)
    doc.add_paragraph(
        '用户服务层负责处理与用户相关的所有业务逻辑：\n\n'
        '（1）认证授权\n'
        '• 用户注册：/api/users/register\n'
        '• 用户登录：/api/users/login（返回JWT Token）\n'
        '• Token刷新：自动续期\n'
        '• 密码修改：/api/users/change-password\n\n'
        '（2）用户信息管理\n'
        '• 获取个人资料：/api/users/profile\n'
        '• 更新个人资料：/api/users/profile (PUT)\n'
        '• 头像上传：支持自定义头像\n\n'
        '（3）用户列表查询\n'
        '• 获取老人列表：/api/users/elder/list\n'
        '• 获取护理人员列表：/api/users/workers\n'
        '• 家属绑定老人：/api/users/binding-elder'
    )
    
    doc.add_heading('7.3 业务逻辑层', level=2)
    doc.add_paragraph(
        '业务逻辑层实现系统核心业务规则：\n\n'
        '（1）护理记录业务\n'
        '• 护理记录创建需要护理人员或管理员权限\n'
        '• 支持按老人ID、护理类型筛选\n'
        '• 记录开始时间和结束时间\n'
        '• 支持状态更新（进行中/已完成）\n\n'
        '（2）健康指标业务\n'
        '• 支持10种健康指标类型\n'
        '• 自动判断指标是否异常\n'
        '• 支持批量导入历史数据\n'
        '• 记录测量人员和测量时间\n\n'
        '（3）护理计划业务\n'
        '• 护理计划包含多个护理任务\n'
        '• 支持周期性任务设置\n'
        '• 任务完成时记录完成时间和执行人\n'
        '• 支持今日任务查询\n\n'
        '（4）订单业务\n'
        '• 根据用户类型限制可见订单\n'
        '• 订单状态流转控制\n'
        '• 订单金额计算（原价-优惠=实付）\n'
        '• 地址信息快照保存'
    )
    
    doc.add_heading('7.4 数据服务层', level=2)
    doc.add_paragraph(
        '数据服务层负责数据持久化和访问：\n\n'
        '（1）数据库连接管理\n'
        '• Flask-SQLAlchemy连接池\n'
        '• 自动提交/回滚事务\n'
        '• 数据库迁移支持\n\n'
        '（2）数据模型定义\n'
        '• User模型：用户信息\n'
        '• HealthMetric模型：健康指标\n'
        '• NursingRecord模型：护理记录\n'
        '• CarePlan模型：护理计划\n'
        '• CareTask模型：护理任务\n'
        '• Order模型：服务订单\n'
        '• Service模型：服务项目\n'
        '• Notification模型：消息通知\n\n'
        '（3）数据查询优化\n'
        '• 索引优化：username, phone, user_type\n'
        '• 分页查询：支持page和page_size参数\n'
        '• 关系查询：使用SQLAlchemy的relationship'
    )
    
    doc.add_page_break()
    
    # ========== 8. 部署视图 ==========
    doc.add_heading('8. 部署视图', level=1)
    doc.add_paragraph(
        '系统支持两种部署模式：\n\n'
        '（1）开发模式部署\n'
        '• 后端：python run.py（Flask开发服务器）\n'
        '  - 运行端口：5000\n'
        '  - 调试模式：开启\n'
        '• 前端：npm run dev（Vite开发服务器）\n'
        '  - 运行端口：5173\n'
        '  - 代理配置：/api -> localhost:5000\n\n'
        '（2）生产模式部署（推荐）\n'
        '• 前端：npm run build\n'
        '  - 构建产物：dist/\n'
        '  - 部署：Nginx静态托管\n'
        '• 后端：Gunicorn + Nginx\n'
        '  - Gunicorn：WSGI服务器\n'
        '  - Nginx：反向代理和静态文件服务\n'
        '• 数据库：SQLite（轻量）或 PostgreSQL（生产）\n\n'
        '（3）环境配置\n'
        '• .env文件配置\n'
        '  - FLASK_ENV=production\n'
        '  - SECRET_KEY=xxx\n'
        '  - XFYUN_API_KEY=xxx（AI服务）\n'
        '  - REDIS配置（如使用Redis）'
    )
    
    # ========== 9. 接口设计 ==========
    doc.add_heading('9. 接口设计', level=1)
    doc.add_paragraph('系统采用RESTful API设计规范：\n')
    doc.add_paragraph(
        'Base URL: /api\n'
        '认证方式: Bearer Token（JWT）\n'
        '响应格式: JSON\n'
        '统一响应结构: {code, message, data}'
    )
    
    doc.add_paragraph('\n主要API接口列表：')
    
    api_table = doc.add_table(rows=16, cols=4)
    api_table.style = 'Table Grid'
    api_data = [
        ['模块', '接口路径', '方法', '说明'],
        ['用户', '/api/users/login', 'POST', '用户登录'],
        ['用户', '/api/users/register', 'POST', '用户注册'],
        ['用户', '/api/users/profile', 'GET/PUT', '用户资料'],
        ['用户', '/api/users/elder/list', 'GET', '老人列表'],
        ['护理', '/api/nursing/records', 'GET/POST', '护理记录'],
        ['健康', '/api/health/metrics', 'GET/POST', '健康指标'],
        ['护理计划', '/api/care/plans', 'GET/POST', '护理计划'],
        ['护理任务', '/api/care/tasks', 'POST', '护理任务'],
        ['订单', '/api/orders/', 'GET/POST', '订单管理'],
        ['服务', '/api/services/', 'GET', '服务列表'],
        ['大数据', '/api/bigdata/health-analysis', 'GET', '健康分析'],
        ['大数据', '/api/bigdata/risk-assessment', 'GET', '风险评估'],
        ['AI助手', '/api/ai/chat', 'POST', 'AI对话'],
        ['统计', '/api/statistics/dashboard', 'GET', '仪表盘']
    ]
    for i, row_data in enumerate(api_data):
        for j, cell_text in enumerate(row_data):
            cell = api_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ========== 10. 系统的组织结构 ==========
    doc.add_heading('10. 系统的组织结构', level=1)
    doc.add_paragraph('后端项目结构：')
    doc.add_paragraph(
        'backend/\n'
        '├── app/\n'
        '│   ├── __init__.py          # 应用工厂\n'
        '│   ├── config.py             # 配置管理\n'
        '│   ├── extensions.py          # 扩展初始化\n'
        '│   ├── models/               # 数据模型\n'
        '│   │   ├── user.py\n'
        '│   │   ├── health.py\n'
        '│   │   ├── nursing.py\n'
        '│   │   ├── care.py\n'
        '│   │   ├── order.py\n'
        '│   │   └── service.py\n'
        '│   ├── routes/               # 路由蓝图\n'
        '│   │   ├── users.py\n'
        '│   │   ├── health.py\n'
        '│   │   ├── nursing.py\n'
        '│   │   ├── care.py\n'
        '│   │   ├── bigdata.py\n'
        '│   │   ├── ai.py\n'
        '│   │   └── order.py\n'
        '│   └── utils/                # 工具函数\n'
        '│       ├── response.py\n'
        '│       └── auth.py\n'
        '├── run.py                     # 入口文件\n'
        '└── requirements.txt          # 依赖清单'
    )
    
    doc.add_paragraph('\n前端项目结构：')
    doc.add_paragraph(
        'frontend/src/\n'
        '├── main.ts                  # 入口文件\n'
        '├── App.vue                  # 根组件\n'
        '├── router/                  # 路由配置\n'
        '│   └── index.ts\n'
        '├── store/                   # 状态管理\n'
        '│   └── auth.ts\n'
        '├── views/                   # 页面组件\n'
        '│   ├── Login.vue\n'
        '│   ├── Dashboard.vue\n'
        '│   ├── BigDataAnalytics.vue\n'
        '│   └── ...\n'
        '├── components/              # 公共组件\n'
        '├── layouts/                 # 布局组件\n'
        '├── styles/                  # 样式文件\n'
        '├── types/                   # 类型定义\n'
        '└── utils/                   # 工具函数'
    )
    
    # ========== 11. 界面设计要求 ==========
    doc.add_heading('11. 界面设计要求', level=1)
    doc.add_paragraph(
        '（1）整体风格\n'
        '• 采用现代简约设计风格，界面清晰美观\n'
        '• 色彩搭配协调，主色调蓝色系\n'
        '• 响应式布局，支持PC和移动端\n\n'
        '（2）导航设计\n'
        '• 左侧导航栏，支持菜单折叠\n'
        '• 顶部显示用户信息和快捷操作\n'
        '• 面包屑导航，显示当前位置\n\n'
        '（3）数据展示\n'
        '• 表格采用斑马线样式\n'
        '• 支持分页展示，每页条数可选\n'
        '• 图表采用ECharts，支持交互\n\n'
        '（4）表单设计\n'
        '• 表单项带必填标识\n'
        '• 输入框带placeholder提示\n'
        '• 支持表单验证和错误提示\n\n'
        '（5）用户体验\n'
        '• 操作按钮位置统一\n'
        '• 支持键盘快捷键\n'
        '• 加载状态有loading动画\n'
        '• 操作成功/失败有提示'
    )
    
    doc.add_page_break()
    
    # ========== 12. 详细设计 ==========
    doc.add_heading('12. 详细设计', level=1)
    
    doc.add_heading('12.1 系统模块划分', level=2)
    
    module_table = doc.add_table(rows=9, cols=3)
    module_table.style = 'Table Grid'
    module_data = [
        ['模块名称', '功能描述', '主要API'],
        ['用户管理', '用户注册、登录、信息管理', '/api/users/*'],
        ['护理记录', '护理服务记录和查询', '/api/nursing/*'],
        ['健康管理', '健康指标监测和管理', '/api/health/*'],
        ['护理计划', '护理计划制定和任务管理', '/api/care/*'],
        ['订单管理', '服务订单全流程管理', '/api/orders/*'],
        ['数据分析', '健康数据统计和分析', '/api/bigdata/*'],
        ['AI助手', '智能健康咨询', '/api/ai/*'],
        ['服务管理', '服务项目管理', '/api/services/*']
    ]
    for i, row_data in enumerate(module_data):
        for j, cell_text in enumerate(row_data):
            cell = module_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_heading('12.2 模块详细设计', level=2)
    doc.add_paragraph('详见各模块设计说明（12.6-12.9节）')
    
    doc.add_heading('12.3 接口设计', level=2)
    doc.add_paragraph('详见第9节接口设计')
    
    doc.add_heading('12.4 算法与模型设计', level=2)
    doc.add_paragraph(
        '（1）统计分析算法\n'
        '• 均值计算：statistics.mean()\n'
        '• 中位数计算：statistics.median()\n'
        '• 标准差计算：statistics.stdev()\n'
        '• 方差计算：statistics.variance()\n\n'
        '（2）异常检测算法\n'
        '• 基于标准范围判断：value < min OR value > max\n'
        '• 健康指标标准范围定义在代码中\n'
        '• 支持自定义标准范围\n\n'
        '（3）趋势预测算法\n'
        '• 使用简单线性回归\n'
        '• 基于历史数据预测未来值\n'
        '• 计算预测斜率判断趋势方向\n\n'
        '（4）风险评估算法\n'
        '• 高风险：异常率 > 50%\n'
        '• 中风险：异常率 > 20%\n'
        '• 低风险：异常率 <= 20%'
    )
    
    doc.add_heading('12.5 系统部署与运维设计', level=2)
    doc.add_paragraph(
        '（1）部署流程\n'
        '• 前端构建：npm run build\n'
        '• 后端部署：Gunicorn启动\n'
        '• Nginx配置反向代理\n'
        '• 数据库初始化和迁移\n\n'
        '（2）监控告警\n'
        '• 服务健康检查接口\n'
        '• 日志记录和查询\n'
        '• 异常告警通知\n\n'
        '（3）备份恢复\n'
        '• 数据库定期备份\n'
        '• 配置文件版本管理\n'
        '• 灾难恢复预案'
    )
    
    doc.add_page_break()
    
    doc.add_heading('12.6 登入注册界面的设计说明', level=2)
    doc.add_paragraph(
        '（1）登录页面设计\n'
        '• 表单字段：用户名/手机号、密码\n'
        '• 登录按钮：蓝色主色调\n'
        '• 记住登录状态选项\n'
        '• 忘记密码链接\n\n'
        '（2）注册页面设计\n'
        '• 表单字段：用户名、密码、确认密码、手机号、验证码\n'
        '• 用户类型选择：老人/护理人员/家属\n'
        '• 服务协议勾选\n'
        '• 注册按钮\n\n'
        '（3）代码实现\n'
        '前端组件：frontend/src/views/Login.vue\n'
        '前端组件：frontend/src/views/Register.vue\n'
        '后端API：\n'
        '• POST /api/users/login - 登录\n'
        '• POST /api/users/register - 注册\n'
        '• POST /api/users/send-sms - 发送短信验证码\n'
        '• POST /api/users/send-email-code - 发送邮箱验证码\n\n'
        '（4）认证流程\n'
        '1. 用户输入用户名密码\n'
        '2. 前端发送登录请求\n'
        '3. 后端验证用户名密码\n'
        '4. 验证成功生成JWT Token\n'
        '5. 前端存储Token到localStorage\n'
        '6. 后续请求携带Token\n'
        '7. 后端验证Token有效性'
    )
    
    doc.add_heading('12.7 护理计划的设计说明', level=2)
    doc.add_paragraph(
        '（1）护理计划模型设计\n\n'
        'CarePlan表结构：\n'
        '• id: 主键\n'
        '• elder_id: 老人ID\n'
        '• title: 计划标题\n'
        '• description: 计划描述\n'
        '• start_date: 开始日期\n'
        '• end_date: 结束日期\n'
        '• status: 状态（1-执行中, 2-已完成, 3-已暂停）\n'
        '• created_by: 创建人ID\n'
        '• created_at: 创建时间\n\n'
        'CareTask表结构：\n'
        '• id: 主键\n'
        '• care_plan_id: 护理计划ID\n'
        '• task_name: 任务名称\n'
        '• task_type: 任务类型\n'
        '• description: 任务描述\n'
        '• frequency: 执行频率\n'
        '• scheduled_time: 计划时间\n'
        '• status: 状态（1-待执行, 2-已完成, 3-已取消）\n'
        '• completed_at: 完成时间\n'
        '• completed_by: 完成人ID\n\n'
        '（2）API接口设计\n'
        '• POST /api/care/plans - 创建护理计划\n'
        '• GET /api/care/plans - 获取计划列表\n'
        '• GET /api/care/plans/<id> - 获取计划详情（含任务）\n'
        '• PUT /api/care/plans/<id> - 更新护理计划\n'
        '• POST /api/care/tasks - 创建护理任务\n'
        '• POST /api/care/tasks/<id>/complete - 完成任务\n'
        '• GET /api/care/tasks/today - 获取今日任务\n\n'
        '（3）业务流程\n'
        '1. 管理员创建护理计划\n'
        '2. 为计划添加护理任务\n'
        '3. 设置任务执行时间和频率\n'
        '4. 护理人员执行任务并记录\n'
        '5. 任务完成后更新状态\n'
        '6. 系统统计任务完成情况'
    )
    
    doc.add_page_break()
    
    doc.add_heading('12.8 健康数据的设计说明', level=2)
    doc.add_paragraph(
        '（1）健康指标模型设计\n\n'
        'HealthMetric表结构：\n'
        '• id: 主键\n'
        '• elder_id: 老人ID\n'
        '• metric_type: 指标类型（1-10）\n'
        '• metric_value: 指标值\n'
        '• unit: 单位\n'
        '• recorded_by: 记录人ID\n'
        '• recorded_at: 记录时间\n'
        '• notes: 备注\n'
        '• created_at: 创建时间\n\n'
        '指标类型定义：\n'
        '• 1-体温: 正常范围36.0-37.3°C\n'
        '• 2-血压(收缩压): 无固定范围\n'
        '• 3-血压(舒张压): 无固定范围\n'
        '• 4-心率: 正常范围60-100 BPM\n'
        '• 5-血氧: 正常范围95-100%\n'
        '• 6-血糖: 正常范围3.9-6.1 mmol/L\n'
        '• 7-体重: 无固定范围\n'
        '• 8-身高: 无固定范围\n'
        '• 9-睡眠时长: 无固定范围\n'
        '• 10-今日步数: 无固定范围\n\n'
        '（2）API接口设计\n'
        '• POST /api/health/metrics - 创建健康记录\n'
        '• GET /api/health/metrics - 获取记录列表\n'
        '• GET /api/health/metrics/<id> - 获取记录详情\n'
        '• GET /api/health/metrics/latest/<elder_id> - 获取最新指标\n'
        '• GET /api/health/metrics/trend - 获取趋势数据\n\n'
        '（3）统计计算\n'
        '• 按指标类型分组统计\n'
        '• 计算：最小值、最大值、平均值、中位数、标准差\n'
        '• 异常率 = 异常记录数 / 总记录数 × 100%\n'
        '• 趋势 = 近期7天均值 - 之前均值'
    )
    
    doc.add_page_break()
    
    doc.add_heading('12.9 大数据分析的设计说明', level=2)
    doc.add_paragraph(
        '（1）大数据分析模块架构\n\n'
        '模块位置：backend/app/routes/bigdata.py\n'
        '蓝图前缀：/api/bigdata\n\n'
        '（2）核心API接口\n\n'
        '1. /api/bigdata/health-analysis (GET)\n'
        '   功能：综合健康分析\n'
        '   参数：elder_id（可选）, days（默认30天）\n'
        '   返回：各指标统计数据（min, max, avg, median, std, abnormal_rate, trend）\n\n'
        '2. /api/bigdata/health-insights (GET)\n'
        '   功能：智能洞察生成\n'
        '   返回：数据完整度、测量频率、趋势分析、异常检测统计\n\n'
        '3. /api/bigdata/anomaly-alerts (GET)\n'
        '   功能：异常告警列表\n'
        '   返回：异常记录列表（时间、类型、值、严重程度）\n\n'
        '4. /api/bigdata/risk-assessment (GET)\n'
        '   功能：老人风险评估\n'
        '   返回：风险评分、风险等级、风险因素\n\n'
        '5. /api/bigdata/nursing-efficiency (GET)\n'
        '   功能：护理效率统计\n'
        '   返回：各类护理记录数、完成率、平均时长\n\n'
        '6. /api/bigdata/care-plan-analysis (GET)\n'
        '   功能：护理计划分析\n'
        '   返回：计划执行情况、任务完成率\n\n'
        '7. /api/bigdata/activity-patterns (GET)\n'
        '   功能：活动模式分析\n'
        '   返回：睡眠时长、步数等日常活动数据\n\n'
        '8. /api/bigdata/comparison (GET)\n'
        '   功能：数据对比分析\n'
        '   返回：老人间健康指标对比\n\n'
        '9. /api/bigdata/health-prediction (GET)\n'
        '   功能：健康趋势预测\n'
        '   参数：elder_id, metric_type, days, predict_days\n'
        '   返回：历史数据、预测数据、置信度\n\n'
        '（3）统计分析算法\n'
        '• 基础统计：均值、中位数、众数、标准差、方差\n'
        '• 异常检测：与健康标准范围对比\n'
        '• 趋势分析：线性回归斜率计算\n'
        '• 风险评估：基于异常率的加权评分\n\n'
        '（4）数据安全标准范围\n'
        '• 体温：36.0-37.3°C\n'
        '• 心率：60-100 BPM\n'
        '• 血氧：95-100%\n'
        '• 血糖：3.9-6.1 mmol/L'
    )
    
    doc.add_page_break()
    
    # ========== 13. 模块相互关系表 ==========
    doc.add_heading('13. 模块相互关系表', level=1)
    
    relation_table = doc.add_table(rows=10, cols=6)
    relation_table.style = 'Table Grid'
    relation_data = [
        ['模块', '用户', '护理', '健康', '订单', '护理计划'],
        ['用户', '-', 'N:1', 'N:1', '1:N', '1:N'],
        ['护理', 'N:1', '-', 'N:1', '-', 'N:1'],
        ['健康', 'N:1', 'N:1', '-', '-', '-'],
        ['订单', '1:N', 'N:1', '-', '-', '-'],
        ['护理计划', 'N:1', '-', '-', '-', '-'],
        ['AI助手', '-', '-', '1:1', '-', '-'],
        ['服务', '-', '1:N', '-', '1:N', '-'],
        ['通知', 'N:1', '-', '-', '-', '-'],
        ['统计', '聚合', '聚合', '聚合', '聚合', '聚合']
    ]
    for i, row_data in enumerate(relation_data):
        for j, cell_text in enumerate(row_data):
            cell = relation_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0 or j == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(9)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph('\n关系说明：')
    doc.add_paragraph(
        '• N:1 - 多对一关系（如多个订单属于一个用户）\n'
        '• 1:N - 一对多关系（如一个用户有多个护理记录）\n'
        '• 1:1 - 一对一关系（如AI配置与系统配置）\n'
        '• 聚合 - 统计模块聚合各业务数据'
    )
    
    # ========== 文档结束 ==========
    doc.add_page_break()
    ending = doc.add_paragraph('— 文档结束 —')
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def main():
    """生成文档"""
    doc = create_document()
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 
                               '架构设计与详细设计说明书.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()