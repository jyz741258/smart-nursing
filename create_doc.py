# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_doc():
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('互联网+智慧护理\n移动护理平台APP')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()
    
    # 文档类型
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run('数据存储设计说明书')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    for _ in range(4):
        doc.add_paragraph()
    
    # 版本信息
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version: 1.1')
    run.font.size = Pt(16)
    
    for _ in range(6):
        doc.add_paragraph()
    
    # 项目信息表格
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('项目承担部门：', '坤坤快跑技术有限公司'),
        ('撰写人（签名）：', '李鹏辉'),
        ('完成日期：', '2026-4-10'),
        ('本文档使用部门：', '■项目组  ■维护人员  ■用户'),
        ('评审负责人（签名）：', '江雨泽'),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
    
    # 评审日期
    review_date = doc.add_paragraph()
    review_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = review_date.add_run('评审日期：2026-4-10')
    run.font.size = Pt(14)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 文档信息页 ====================
    doc.add_paragraph()
    heading = doc.add_paragraph()
    run = heading.add_run('文档信息')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc_info_table = doc.add_table(rows=5, cols=2)
    doc_info_data = [
        ('标题：', '互联网+智慧护理移动护理平台APP数据存储设计说明书'),
        ('作者：', '李鹏辉'),
        ('创建日期：', '2026-4-10'),
        ('上次更新日期：', '2026-4-10'),
        ('版本：', '1.0.0'),
    ]
    for i, (label, value) in enumerate(doc_info_data):
        doc_info_table.rows[i].cells[0].text = label
        doc_info_table.rows[i].cells[1].text = value
        doc_info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 修订历史
    doc.add_paragraph()
    heading2 = doc.add_paragraph()
    run = heading2.add_run('修订文档历史记录')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    revision_table = doc.add_table(rows=4, cols=4)
    revision_table.style = 'Table Grid'
    revision_headers = ['日期', '版本', '说明', '作者']
    for i, header in enumerate(revision_headers):
        revision_table.rows[0].cells[i].text = header
        revision_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    revision_data = [
        ('2026-4-10', '1.0.0', '初始版本，完成数据库基础架构设计', '李鹏辉'),
        ('2026-4-10', '1.1.0', '完善数据表结构说明，补充实体关系图', '李鹏辉'),
        ('', '', '', ''),
    ]
    for i, row_data in enumerate(revision_data):
        for j, cell_data in enumerate(row_data):
            revision_table.rows[i+1].cells[j].text = cell_data
    
    # 分页
    doc.add_page_break()
    
    # ==================== 目录占位 ====================
    doc.add_paragraph()
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run('目 录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    toc_items = [
        '1. 引言',
        '1.1 编写目的',
        '1.2 背景',
        '1.3 定义',
        '1.4 参考资料',
        '2. 数据存储架构设计',
        '2.1 数据存储技术选型',
        '2.2 数据存储部署视图',
        '3. 数据存储详细设计',
        '3.1 结构化数据库',
        '3.2 非结构化数据库',
        '3.3 列式数据库',
        '3.4 分布式文件系统存储',
        '4. 结构设计',
        '4.1 结构化数据库',
        '4.2 非结构化数据库',
        '4.3 列式数据库',
        '4.4 分布式文件系统存储',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    # 分页
    doc.add_page_break()
    
    # ==================== 1. 引言 ====================
    h1 = doc.add_heading('1. 引言', level=1)
    
    h2 = doc.add_heading('1.1 编写目的', level=2)
    doc.add_paragraph(
        '数据库设计说明书是根据概要设计说明书的要求所编写的，是为详细设计作依据的，'
        '为详细设计提供标准，并提供给编码人员和测试人员。'
    )
    
    h2 = doc.add_heading('1.2 背景', level=2)
    bg_text = (
        '使用此数据库的软件系统的名称：KKhealth 互联网+智慧护理移动护理平台APP\n'
        '该软件系统开发项目的任务提出者：李鹏辉、陈问未\n'
        '该软件系统的用户：实训学员、教师、平台管理员'
    )
    doc.add_paragraph(bg_text)
    
    h2 = doc.add_heading('1.3 定义', level=2)
    def_table = doc.add_table(rows=4, cols=2)
    def_table.style = 'Table Grid'
    def_data = [
        ('KKhealth', '互联网+智慧护理移动护理平台APP'),
        ('CDM', '概念数据模型（Concept Data Model）'),
        ('PDM', '物理数据模型（Physics Data Model）'),
        ('MySQL', '关系型数据库管理系统'),
    ]
    for i, (term, meaning) in enumerate(def_data):
        def_table.rows[i].cells[0].text = term
        def_table.rows[i].cells[1].text = meaning
        def_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('1.4 参考资料', level=2)
    doc.add_paragraph('《互联网+智慧护理移动护理平台APP》软件需求规约')
    
    # 分页
    doc.add_page_break()
    
    # ==================== 2. 数据存储架构设计 ====================
    h1 = doc.add_heading('2. 数据存储架构设计', level=1)
    
    h2 = doc.add_heading('2.1 数据存储技术选型', level=2)
    
    tech_items = [
        ('结构化数据库', '如MySQL、PostgreSQL等，适用于存储关系型数据，支持事务处理、复杂查询等操作。'),
        ('非结构化数据库', '如MongoDB、Cassandra等NoSQL数据库，适用于存储非关系型数据，如文档、键值对等，具有高度的灵活性和可扩展性。'),
        ('列式数据库', '如HBase、Amazon Redshift等，适用于存储和分析海量数据，特别适用于OLAP（在线分析处理）场景。'),
        ('分布式文件系统', '如Hadoop HDFS、Ceph等，适用于存储大规模非结构化数据，如日志文件、图片、视频等。'),
    ]
    
    for title, desc in tech_items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}：')
        run.font.bold = True
        p.add_run(desc)
    
    h2 = doc.add_heading('2.2 数据存储部署视图', level=2)
    doc.add_paragraph('（详见物理模型部分）')
    
    # 分页
    doc.add_page_break()
    
    # ==================== 3. 数据存储详细设计 ====================
    h1 = doc.add_heading('3. 数据存储详细设计', level=1)
    
    h2 = doc.add_heading('3.1 结构化数据库', level=2)
    
    h3 = doc.add_heading('3.1.1 类型划分', level=3)
    doc.add_paragraph('主表：7个')
    
    h3 = doc.add_heading('3.1.2 标识符和约定', level=3)
    doc.add_paragraph('数据库表的命名采用表名的英文或英文缩写，便于编程实现。')
    
    id_table = doc.add_table(rows=5, cols=3)
    id_table.style = 'Table Grid'
    id_headers = ['序号', '表名', '描述']
    for i, h in enumerate(id_headers):
        id_table.rows[0].cells[i].text = h
        id_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    id_data = [
        ('1', 'sm_users', '用户信息表'),
        ('2', 'sm_elder', '老人信息表'),
        ('3', 'sm_nursing_record', '护理记录表'),
        ('4', 'sm_health_data', '健康数据表'),
    ]
    for i, row in enumerate(id_data):
        for j, val in enumerate(row):
            id_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('3.1.3 数据库表设计', level=3)
    
    # 表3.1.3.1 用户表
    doc.add_paragraph('3.1.3.1 用户信息表（sm_users）')
    users_table = doc.add_table(rows=9, cols=5)
    users_table.style = 'Table Grid'
    users_headers = ['字段中文名', '字段名', '字段类型', '主键', '说明']
    for i, h in enumerate(users_headers):
        users_table.rows[0].cells[i].text = h
        users_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    users_data = [
        ('用户ID', 'id', 'bigint', '是', '自增主键'),
        ('用户名', 'username', 'varchar(50)', '否', '登录用户名'),
        ('密码', 'password', 'varchar(255)', '否', '加密存储'),
        ('用户类型', 'user_type', 'int', '否', '1=老人 2=护理员 3=管理员 4=家属'),
        ('姓名', 'name', 'varchar(50)', '否', '真实姓名'),
        ('手机号', 'phone', 'varchar(11)', '否', '唯一'),
        ('创建时间', 'create_time', 'datetime', '否', ''),
        ('更新时间', 'update_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(users_data):
        for j, val in enumerate(row):
            users_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.2 老人信息表
    doc.add_paragraph('3.1.3.2 老人信息表（sm_elder）')
    elder_table = doc.add_table(rows=12, cols=5)
    elder_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        elder_table.rows[0].cells[i].text = h
        elder_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    elder_data = [
        ('老人ID', 'id', 'bigint', '是', '自增主键'),
        ('用户ID', 'user_id', 'bigint', '否', '外键关联用户表'),
        ('姓名', 'name', 'varchar(50)', '否', ''),
        ('性别', 'gender', 'varchar(10)', '否', '男/女'),
        ('年龄', 'age', 'int', '否', ''),
        ('身份证号', 'id_card', 'varchar(18)', '否', '唯一'),
        ('健康状况', 'health_status', 'varchar(200)', '否', ''),
        ('护理等级', 'care_level', 'int', '否', '1=一级 2=二级 3=三级'),
        ('紧急联系人', 'emergency_contact', 'varchar(50)', '否', ''),
        ('紧急电话', 'emergency_phone', 'varchar(11)', '否', ''),
        ('创建时间', 'create_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(elder_data):
        for j, val in enumerate(row):
            elder_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.3 护理记录表
    doc.add_paragraph('3.1.3.3 护理记录表（sm_nursing_record）')
    nursing_table = doc.add_table(rows=10, cols=5)
    nursing_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        nursing_table.rows[0].cells[i].text = h
        nursing_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    nursing_data = [
        ('记录ID', 'id', 'bigint', '是', '自增主键'),
        ('老人ID', 'elder_id', 'bigint', '否', '外键关联老人表'),
        ('护理员ID', 'nurse_id', 'bigint', '否', '外键关联用户表'),
        ('护理项目', 'item', 'varchar(100)', '否', ''),
        ('护理内容', 'content', 'text', '否', ''),
        ('执行时间', 'execute_time', 'datetime', '否', ''),
        ('状态', 'status', 'int', '否', '1=待执行 2=执行中 3=已完成'),
        ('备注', 'remark', 'varchar(500)', '否', ''),
        ('创建时间', 'create_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(nursing_data):
        for j, val in enumerate(row):
            nursing_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.4 健康数据表
    doc.add_paragraph('3.1.3.4 健康数据表（sm_health_data）')
    health_table = doc.add_table(rows=11, cols=5)
    health_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        health_table.rows[0].cells[i].text = h
        health_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    health_data = [
        ('数据ID', 'id', 'bigint', '是', '自增主键'),
        ('老人ID', 'elder_id', 'bigint', '否', '外键关联老人表'),
        ('血压', 'blood_pressure', 'varchar(20)', '否', '如：120/80'),
        ('心率', 'heart_rate', 'int', '否', '单位：次/分'),
        ('血糖', 'blood_sugar', 'decimal(5,2)', '否', '单位：mmol/L'),
        ('体温', 'temperature', 'decimal(4,1)', '否', '单位：℃'),
        ('记录时间', 'record_time', 'datetime', '否', ''),
        ('记录人', 'recorder_id', 'bigint', '否', '外键关联用户表'),
        ('备注', 'remark', 'varchar(500)', '否', ''),
        ('创建时间', 'create_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(health_data):
        for j, val in enumerate(row):
            health_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.5 护理计划表
    doc.add_paragraph('3.1.3.5 护理计划表（sm_care_plan）')
    plan_table = doc.add_table(rows=9, cols=5)
    plan_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        plan_table.rows[0].cells[i].text = h
        plan_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    plan_data = [
        ('计划ID', 'id', 'bigint', '是', '自增主键'),
        ('老人ID', 'elder_id', 'bigint', '否', '外键关联老人表'),
        ('计划名称', 'plan_name', 'varchar(100)', '否', ''),
        ('计划内容', 'content', 'text', '否', ''),
        ('开始日期', 'start_date', 'date', '否', ''),
        ('结束日期', 'end_date', 'date', '否', ''),
        ('状态', 'status', 'int', '否', '1=待执行 2=执行中 3=已完成 4=已取消'),
        ('创建时间', 'create_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(plan_data):
        for j, val in enumerate(row):
            plan_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.6 通知公告表
    doc.add_paragraph('3.1.3.6 通知公告表（sm_notifications）')
    notif_table = doc.add_table(rows=9, cols=5)
    notif_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        notif_table.rows[0].cells[i].text = h
        notif_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    notif_data = [
        ('通知ID', 'id', 'bigint', '是', '自增主键'),
        ('标题', 'title', 'varchar(200)', '否', ''),
        ('内容', 'content', 'text', '否', ''),
        ('类型', 'type', 'int', '否', '1=系统通知 2=健康提醒 3=活动公告'),
        ('发送者', 'sender_id', 'bigint', '否', '外键关联用户表'),
        ('目标用户类型', 'target_type', 'int', '否', '0=全部 1=老人 2=护理员等'),
        ('已读状态', 'is_read', 'int', '否', '0=未读 1=已读'),
        ('创建时间', 'create_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(notif_data):
        for j, val in enumerate(row):
            notif_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # 表3.1.3.7 服务订单表
    doc.add_paragraph('3.1.3.7 服务订单表（sm_service_order）')
    order_table = doc.add_table(rows=11, cols=5)
    order_table.style = 'Table Grid'
    for i, h in enumerate(users_headers):
        order_table.rows[0].cells[i].text = h
        order_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    order_data = [
        ('订单ID', 'id', 'bigint', '是', '自增主键'),
        ('订单编号', 'order_no', 'varchar(50)', '否', '唯一'),
        ('老人ID', 'elder_id', 'bigint', '否', '外键关联老人表'),
        ('家属ID', 'family_id', 'bigint', '否', '外键关联用户表'),
        ('服务类型', 'service_type', 'varchar(50)', '否', ''),
        ('服务时间', 'service_time', 'datetime', '否', ''),
        ('金额', 'amount', 'decimal(10,2)', '否', '单位：元'),
        ('订单状态', 'status', 'int', '否', '1=待支付 2=已支付 3=服务中 4=已完成'),
        ('创建时间', 'create_time', 'datetime', '否', ''),
        ('更新时间', 'update_time', 'datetime', '否', ''),
    ]
    for i, row in enumerate(order_data):
        for j, val in enumerate(row):
            order_table.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('3.1.4 支持软件', level=3)
    doc.add_paragraph('• MySQL Server 5.7 或更高版本\n• Navicat Premium 15（数据库管理工具）\n• PowerDesigner 16.7（可选，数据库建模工具）')
    
    doc.add_page_break()
    
    # ==================== 3.2-3.4 ====================
    h2 = doc.add_heading('3.2 非结构化数据库', level=2)
    doc.add_paragraph('本系统目前主要使用关系型数据库存储核心业务数据，暂不涉及非结构化数据存储需求。如后续有图片、日志文件等存储需求，可考虑引入MongoDB或OSS对象存储服务。')
    
    h2 = doc.add_heading('3.3 列式数据库', level=2)
    doc.add_paragraph('本系统当前数据量级暂不需要列式数据库。如后续需要进行大规模数据分析和报表统计，可考虑引入ClickHouse或Apache Doris等OLAP引擎。')
    
    h2 = doc.add_heading('3.4 分布式文件系统存储', level=2)
    doc.add_paragraph('本系统目前暂无大规模文件存储需求。如后续需要存储用户头像、护理记录附件等文件，可考虑使用七牛云OSS或阿里云OSS服务。')
    
    doc.add_page_break()
    
    # ==================== 4. 结构设计 ====================
    h1 = doc.add_heading('4. 结构设计', level=1)
    
    h2 = doc.add_heading('4.1 结构化数据库', level=2)
    
    h3 = doc.add_heading('4.1.1 概念模型（CDM）', level=3)
    doc.add_paragraph('系统包含以下核心实体：')
    entities = [
        '• 用户（User）：系统使用者，包括老人、护理员、管理员、家属',
        '• 老人（Elder）：需要护理服务的老年用户',
        '• 护理记录（NursingRecord）：护理员对老人的护理服务记录',
        '• 健康数据（HealthData）：老人的各项健康指标数据',
        '• 护理计划（CarePlan）：为老人制定的护理计划',
        '• 通知公告（Notification）：系统推送的消息',
        '• 服务订单（ServiceOrder）：家属为老人预约的服务订单',
    ]
    for e in entities:
        doc.add_paragraph(e)
    
    doc.add_paragraph()
    doc.add_paragraph('实体关系说明：')
    relations = [
        '• 一个用户可以对应一个老人信息（一对一）',
        '• 一个老人可以有多个护理记录（一对多）',
        '• 一个老人可以有多个健康数据记录（一对多）',
        '• 一个老人可以有多个护理计划（一对多）',
        '• 一个用户可以创建多个通知公告（一对多）',
        '• 一个老人可以关联多个服务订单（一对多）',
    ]
    for r in relations:
        doc.add_paragraph(r)
    
    h3 = doc.add_heading('4.1.2 物理模型（PDM）', level=3)
    doc.add_paragraph('基于MySQL 5.7实现的物理数据模型，主要包含以下数据表：')
    
    pdm_tables = [
        '• sm_users - 用户信息表',
        '• sm_elder - 老人信息表',
        '• sm_nursing_record - 护理记录表',
        '• sm_health_data - 健康数据表',
        '• sm_care_plan - 护理计划表',
        '• sm_notifications - 通知公告表',
        '• sm_service_order - 服务订单表',
    ]
    for t in pdm_tables:
        doc.add_paragraph(t)
    
    doc.add_paragraph()
    doc.add_paragraph('数据库命名规范：')
    norms = [
        '• 表名统一使用小写字母，单词间用下划线分隔',
        '• 字段名使用小写字母，单词间用下划线分隔',
        '• 主键统一命名为id，类型为bigint自增',
        '• 时间字段使用datetime类型，后缀为_time',
        '• 状态字段使用int类型，前缀表明含义',
    ]
    for n in norms:
        doc.add_paragraph(n)
    
    h2 = doc.add_heading('4.2 非结构化数据库', level=2)
    doc.add_paragraph('暂不涉及。')
    
    h2 = doc.add_heading('4.3 列式数据库', level=2)
    doc.add_paragraph('暂不涉及。')
    
    h2 = doc.add_heading('4.4 分布式文件系统存储', level=2)
    doc.add_paragraph('暂不涉及。')
    
    # 保存文档
    output_path = 'c:/Users/Jyz74/Desktop/数据存储设计说明书.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_doc()
